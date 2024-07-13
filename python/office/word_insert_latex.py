#!/usr/bin/env python3

"""
Insert Latex equations into Word

pip install python-docx latex2mathml mathml2omml

Author: Zihan Chen
Date: 2024-07-13
"""

from docx import Document
from lxml import etree
from latex2mathml.converter import convert
import mathml2omml


def latex_to_omml(eq:str):
    mm = convert(eq)
    om = mathml2omml.convert(mm)
    om_str = om.replace(
        '<m:oMath>',
        '<m:oMath xmlns:m="http://schemas.openxmlformats.org/officeDocument/2006/math">'
    )
    return etree.fromstring(om_str)

def add_omml_to_paragraph(paragraph, omml_element):
    # Append the OMML to the paragraph
    paragraph._element.append(omml_element)

# Create a new document
doc = Document()

# Add a paragraph with text
p = doc.add_paragraph("hello world this is an equation haha ")

omml_element = latex_to_omml('x^3 = y')

# Add the OMML to the paragraph
add_omml_to_paragraph(p, omml_element)


p = doc.add_paragraph("要符合")
add_omml_to_paragraph(p, latex_to_omml('\\overline{x}'))
p.add_run('相关法规')

# Save the document
doc.save("latex_equation.docx")
