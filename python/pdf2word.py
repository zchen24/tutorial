#!/usr/bin/env python3

"""
Put each page of a PDF file as an image into a WORD document
"""

import docx
from docx.shared import Inches
from tkinter import filedialog
from tkinter import Tk
from pdf2image import convert_from_path, convert_from_bytes
import os


root = Tk()
root.withdraw()

my_file = filedialog.askopenfilename()
print('File: {}'.format(my_file))
images = convert_from_path(my_file, dpi=400)

doc = docx.Document()

ratio = 0.68
for i in images:
    f_name = 'tmp.jpg'
    i.save(f_name)
    p = doc.add_paragraph()
    run = p.add_run()
    run.add_break()
    doc.add_picture(f_name,
                    width=Inches(ratio * 8.5))
    doc.add_page_break()
    os.remove(f_name)

doc.save('tmp.docx')
